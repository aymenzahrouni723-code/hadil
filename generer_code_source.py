# -*- coding: utf-8 -*-
"""
Generation du document Word : Langages et Code Source de l'application HydroNFT
Pour l'encadreur Mme. ZAINEB JABRI
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE_DIR, "Code_Source_Application_HydroNFT.docx")

# ========================================
# COULEURS
# ========================================
BLUE_HEADER = RGBColor(0x1F, 0x49, 0x7D)    # Bleu fonce pour les titres
BLUE_LIGHT = 'D6E4F0'                         # Fond bleu clair pour en-tetes tableaux
CODE_BG = 'F2F2F2'                             # Fond gris clair pour le code
GREEN_KEYWORD = RGBColor(0x00, 0x80, 0x00)    # Vert pour les mots-cles
BLUE_STRING = RGBColor(0x00, 0x00, 0xCC)      # Bleu pour les chaines
PURPLE_FUNC = RGBColor(0x80, 0x00, 0x80)      # Violet pour les fonctions
GRAY_COMMENT = RGBColor(0x60, 0x80, 0x60)     # Gris-vert pour les commentaires
BLACK = RGBColor(0, 0, 0)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)


def setup_styles(doc):
    """Configurer les styles du document."""
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for level, size, color in [(1, 16, BLUE_HEADER), (2, 14, BLUE_HEADER), (3, 13, BLUE_HEADER)]:
        s = doc.styles[f'Heading {level}']
        s.font.name = 'Times New Roman'
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        s.paragraph_format.space_after = Pt(8)
        s.paragraph_format.line_spacing = 1.5


def add_p(doc, text, bold=False, italic=False, space_after=6, indent=True, size=12):
    """Ajouter un paragraphe."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    return p


def add_bullet(doc, text, size=12):
    """Ajouter un point de liste."""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)


def add_table(doc, headers, rows, caption=None, header_bg=BLUE_LIGHT):
    """Ajouter un tableau formate."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-tetes
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = cell._element.get_or_add_tcPr()
        s = shading.makeelement(qn('w:shd'), {qn('w:fill'): header_bg, qn('w:val'): 'clear'})
        shading.append(s)

    # Donnees
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

    doc.add_paragraph()


def add_code_block(doc, code_text, language="", filename=""):
    """Ajouter un bloc de code formate avec fond gris."""

    # Titre du fichier
    if filename:
        p = doc.add_paragraph()
        run = p.add_run(f"Fichier : {filename}")
        run.bold = True
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.font.color.rgb = BLUE_HEADER
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(8)

        if language:
            run2 = p.add_run(f"  ({language})")
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(10)
            run2.font.color.rgb = DARK_GRAY
            run2.italic = True

    # Bloc de code dans un tableau a une cellule (pour fond gris)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Largeur du tableau
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = tbl.makeelement(qn('w:tblPr'), {})
        tbl.insert(0, tblPr)

    cell = table.rows[0].cells[0]
    # Fond gris
    shading = cell._element.get_or_add_tcPr()
    s = shading.makeelement(qn('w:shd'), {qn('w:fill'): CODE_BG, qn('w:val'): 'clear'})
    shading.append(s)

    # Nettoyer la cellule
    cell.text = ''

    lines = code_text.strip().split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()

        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15

        # Numero de ligne
        num_run = p.add_run(f"{i+1:3d}  ")
        num_run.font.name = 'Consolas'
        num_run.font.size = Pt(7.5)
        num_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Contenu de la ligne
        code_run = p.add_run(line)
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(7.5)
        code_run.font.color.rgb = BLACK

    doc.add_paragraph()  # Espace apres le bloc


def read_file(filepath):
    """Lire le contenu d'un fichier."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"# Erreur lecture : {e}"


def build_document(doc):
    """Construire le document complet."""

    # =============================================
    # PAGE DE TITRE
    # =============================================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Langages et Code Source\nde l'Application HydroNFT")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = BLUE_HEADER

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Système Hydroponique NFT Intelligent\npour la Culture Urbaine")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    # Infos
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "Réalisé par : BOUDHRIOUWA Hadil",
        "Encadrante : Mme. ZAINEB JABRI",
        "Faculté des Sciences de Tunis",
        "Année Universitaire : 2025/2026"
    ]:
        run = info.add_run(line + "\n")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    doc.add_page_break()

    # =============================================
    # TABLE DES MATIERES
    # =============================================
    doc.add_heading("Table des matières", level=1)

    toc_items = [
        "1. Langages de programmation utilisés",
        "   1.1. Vue d'ensemble des technologies",
        "   1.2. JavaScript (langage principal)",
        "   1.3. Python (backend alternatif)",
        "   1.4. HTML5 et CSS3 (interface)",
        "   1.5. SQL (base de données)",
        "2. Architecture et structure du projet",
        "3. Code source du Frontend (React.js)",
        "   3.1. Point d'entrée (main.jsx)",
        "   3.2. Routeur principal (App.jsx)",
        "   3.3. Tableau de bord (Dashboard.jsx)",
        "   3.4. Composant Jauge (GaugeCard.jsx)",
        "   3.5. Contrôle système (Control.jsx)",
        "   3.6. Historique capteurs (History.jsx)",
        "   3.7. Alertes (Alerts.jsx)",
        "   3.8. Cultures (Cultures.jsx)",
        "   3.9. Barre de navigation (Navbar.jsx)",
        "4. Code source du Backend (Python Flask)",
        "   4.1. Serveur API REST (app.py)",
        "   4.2. Modèles de données (models.py)",
        "   4.3. Simulateur capteurs ESP32 (sensor_simulator.py)",
        "   4.4. Logique de contrôle (control_logic.py)",
        "5. Code source du Design (CSS)",
        "   5.1. Design System complet (index.css)",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # =============================================
    # 1. LANGAGES DE PROGRAMMATION
    # =============================================
    doc.add_heading("1. Langages de programmation utilisés", level=1)

    doc.add_heading("1.1. Vue d'ensemble des technologies", level=2)
    add_p(doc,
        "L'application HydroNFT a été développée en utilisant plusieurs langages de "
        "programmation modernes, chacun ayant un rôle spécifique dans l'architecture "
        "du système. Le choix de ces technologies a été guidé par les besoins du projet : "
        "interface utilisateur réactive, communication en temps réel avec les capteurs, "
        "et gestion efficace des données."
    )

    add_table(doc,
        ["Langage", "Version", "Rôle dans le projet", "Fichiers concernés"],
        [
            ["JavaScript (ES6+)", "ECMAScript 2020+", "Langage principal : frontend React.js + backend Node.js", "*.jsx, *.js"],
            ["Python", "3.x", "Backend alternatif : API Flask + simulateur capteurs", "*.py"],
            ["HTML5", "5", "Structure de la page web", "index.html"],
            ["CSS3", "3", "Design System complet (Dark Mode)", "index.css"],
            ["SQL", "SQLite 3", "Requêtes base de données", "Dans app.py, models.py"],
            ["JSX", "React 19", "Extension JavaScript pour les composants UI", "*.jsx"],
        ],
        caption="Tableau 1 : Langages de programmation utilisés dans l'application HydroNFT"
    )

    # --- JavaScript ---
    doc.add_heading("1.2. JavaScript (ES6+) — Langage principal", level=2)
    add_p(doc,
        "JavaScript est le langage principal de l'application HydroNFT. Il est utilisé "
        "à la fois côté client (frontend) et côté serveur (backend) :"
    )
    add_bullet(doc, "Frontend : React.js 19 avec JSX — framework de composants pour l'interface utilisateur interactive.")
    add_bullet(doc, "Backend : Node.js + Express.js — serveur API REST gérant la logique métier, la simulation des capteurs et la gestion de la base de données.")
    add_bullet(doc, "Visualisation : Chart.js 4.5 — bibliothèque de graphiques pour l'affichage de l'historique des capteurs.")

    add_p(doc,
        "Le choix de JavaScript permet d'avoir un langage unique pour tout le projet (full-stack JavaScript), "
        "ce qui facilite la maintenance et la cohérence du code. La syntaxe ES6+ (ECMAScript 2020+) "
        "offre des fonctionnalités modernes comme les fonctions fléchées (arrow functions), "
        "l'async/await pour les appels API, et les modules import/export.",
        space_after=8
    )

    add_table(doc,
        ["Framework / Bibliothèque", "Version", "Rôle"],
        [
            ["React.js", "19.2.6", "Framework d'interface utilisateur (composants réactifs)"],
            ["React Router DOM", "7.15.1", "Navigation entre les 5 écrans de l'application"],
            ["Chart.js", "4.5.1", "Graphiques interactifs (historique des capteurs)"],
            ["react-chartjs-2", "5.3.1", "Wrapper React pour Chart.js"],
            ["Vite", "8.0.12", "Outil de build et serveur de développement (compilation rapide)"],
            ["Express.js", "4.21.0", "Framework serveur Node.js (API REST)"],
            ["sql.js", "1.11.0", "SQLite compilé en WebAssembly pour Node.js"],
            ["cors", "2.8.5", "Middleware pour les requêtes cross-origin"],
        ],
        caption="Tableau 2 : Frameworks et bibliothèques JavaScript utilisés"
    )

    # --- Python ---
    doc.add_heading("1.3. Python — Backend alternatif", level=2)
    add_p(doc,
        "Python est utilisé pour le backend alternatif de l'application, "
        "développé avec le framework Flask. Ce backend implémente :"
    )
    add_bullet(doc, "Flask : micro-framework web pour exposer l'API REST (12 endpoints).")
    add_bullet(doc, "SQLite3 : module intégré pour la gestion de la base de données.")
    add_bullet(doc, "Simulateur de capteurs : simulation réaliste des données ESP32 (pH, EC, température, niveau d'eau) avec bruit gaussien et dérive.")
    add_bullet(doc, "Logique de contrôle automatique : algorithme de vérification des seuils et déclenchement des alertes.")

    add_table(doc,
        ["Module Python", "Rôle"],
        [
            ["Flask", "Serveur web et API REST"],
            ["Flask-CORS", "Gestion des requêtes cross-origin"],
            ["sqlite3", "Interface avec la base de données SQLite"],
            ["threading", "Exécution en arrière-plan du simulateur et du contrôle"],
            ["random", "Génération de bruit réaliste pour les capteurs"],
            ["datetime", "Horodatage des lectures et alertes"],
        ],
        caption="Tableau 3 : Modules Python utilisés dans le backend"
    )

    # --- HTML/CSS ---
    doc.add_heading("1.4. HTML5 et CSS3 — Interface utilisateur", level=2)
    add_p(doc,
        "HTML5 fournit la structure sémantique de la page web principale. "
        "CSS3 implémente un Design System complet avec :"
    )
    add_bullet(doc, "Variables CSS (Custom Properties) pour les couleurs, espacements et typographie.")
    add_bullet(doc, "Thème Dark Mode avec palette inspirée de la nature (vert #2ECC71, bleu #3498DB, turquoise #1ABC9C).")
    add_bullet(doc, "Animations CSS (transitions, keyframes) pour une interface fluide et moderne.")
    add_bullet(doc, "Design responsive (mobile-first) avec media queries.")
    add_bullet(doc, "Glassmorphism (backdrop-filter: blur) pour les cartes et la navbar.")
    add_bullet(doc, "Polices Google Fonts : Outfit (titres) et Inter (corps de texte).")

    # --- SQL ---
    doc.add_heading("1.5. SQL (SQLite) — Base de données", level=2)
    add_p(doc,
        "Le langage SQL est utilisé pour les opérations sur la base de données SQLite "
        "embarquée dans l'application. La base contient 4 tables :"
    )
    add_bullet(doc, "sensor_readings : stockage des lectures de capteurs (pH, EC, température, niveau d'eau) avec horodatage.")
    add_bullet(doc, "culture_profiles : profils des 3 cultures (Laitue, Basilic, Fraise) avec leurs paramètres optimaux.")
    add_bullet(doc, "alerts : alertes générées par le système (type, sévérité, message, acquittement).")
    add_bullet(doc, "system_state : état courant du système (pompe, éclairage, mode, culture active, seuils).")

    doc.add_page_break()

    # =============================================
    # 2. ARCHITECTURE
    # =============================================
    doc.add_heading("2. Architecture et structure du projet", level=1)
    add_p(doc,
        "L'application suit une architecture client-serveur à trois couches (3-tier) :"
    )

    add_table(doc,
        ["Couche", "Technologie", "Fichiers"],
        [
            ["Présentation (Frontend)", "React.js 19 + Vite + CSS3", "App.jsx, Dashboard.jsx, Control.jsx, History.jsx, Alerts.jsx, Cultures.jsx, GaugeCard.jsx, Navbar.jsx, index.css"],
            ["Logique métier (Backend)", "Python Flask / Node.js Express", "app.py, control_logic.py, sensor_simulator.py / server.js"],
            ["Données", "SQLite", "models.py, hydroponic.db"],
        ],
        caption="Tableau 4 : Architecture 3-tier de l'application"
    )

    add_p(doc, "Arborescence du projet :", bold=True, indent=False, space_after=4)

    tree = """HydroNFT/
├── frontend/                    (Application React.js)
│   ├── src/
│   │   ├── main.jsx            (Point d'entrée React)
│   │   ├── App.jsx             (Routeur principal - 5 routes)
│   │   ├── index.css           (Design System complet - 966 lignes)
│   │   ├── components/
│   │   │   ├── GaugeCard.jsx   (Composant jauge SVG circulaire)
│   │   │   └── Navbar.jsx      (Barre de navigation inférieure)
│   │   └── pages/
│   │       ├── Dashboard.jsx   (Tableau de bord - 4 jauges temps réel)
│   │       ├── Control.jsx     (Contrôle pompe/éclairage/consignes)
│   │       ├── History.jsx     (Graphiques Chart.js + statistiques)
│   │       ├── Alerts.jsx      (Système d'alertes et notifications)
│   │       └── Cultures.jsx    (Profils des 3 cultures)
│   ├── package.json            (Dépendances frontend)
│   └── vite.config.js          (Configuration Vite)
├── backend/                     (Serveur Python Flask)
│   ├── app.py                  (API REST - 12 endpoints - 359 lignes)
│   ├── models.py               (Modèles SQLite - 4 tables - 175 lignes)
│   ├── sensor_simulator.py     (Simulateur ESP32 - 134 lignes)
│   ├── control_logic.py        (Contrôle automatique - 176 lignes)
│   └── hydroponic.db           (Base de données SQLite)
├── server.js                    (Backend Node.js alternatif)
├── package.json                 (Dépendances serveur Node.js)
└── render.yaml                  (Configuration déploiement cloud)"""

    add_code_block(doc, tree, "Arborescence", "Structure du projet HydroNFT")

    doc.add_page_break()

    # =============================================
    # 3. CODE SOURCE FRONTEND
    # =============================================
    doc.add_heading("3. Code source du Frontend (React.js)", level=1)
    add_p(doc,
        "Cette section présente le code source complet des composants React de l'application "
        "mobile HydroNFT. Chaque fichier correspond à un écran ou un composant réutilisable."
    )

    # --- 3.1 main.jsx ---
    doc.add_heading("3.1. Point d'entrée de l'application (main.jsx)", level=2)
    add_p(doc,
        "Ce fichier est le point d'entrée de l'application React. Il monte le composant "
        "racine App dans l'élément DOM #root avec le mode strict de React activé."
    )
    add_code_block(doc,
        read_file(os.path.join(BASE_DIR, "frontend", "src", "main.jsx")),
        "JavaScript (JSX)", "frontend/src/main.jsx"
    )

    # --- 3.2 App.jsx ---
    doc.add_heading("3.2. Routeur principal (App.jsx)", level=2)
    add_p(doc,
        "Le composant App définit le routage de l'application avec 5 routes correspondant "
        "aux 5 écrans principaux. Il vérifie également le nombre d'alertes non acquittées "
        "toutes les 5 secondes pour mettre à jour le badge de la barre de navigation."
    )
    add_code_block(doc,
        read_file(os.path.join(BASE_DIR, "frontend", "src", "App.jsx")),
        "JavaScript (JSX)", "frontend/src/App.jsx"
    )

    # --- 3.3 Dashboard.jsx ---
    doc.add_heading("3.3. Tableau de bord (Dashboard.jsx)", level=2)
    add_p(doc,
        "L'écran principal de l'application affiche les 4 jauges des capteurs en temps réel, "
        "les statuts du système (pompe, éclairage, mode) et la dernière alerte. Les données "
        "sont rafraîchies automatiquement toutes les 3 secondes via l'API REST."
    )
    add_code_block(doc,
        read_file(os.path.join(BASE_DIR, "frontend", "src", "pages", "Dashboard.jsx")),
        "JavaScript (JSX)", "frontend/src/pages/Dashboard.jsx"
    )

    # --- 3.4 GaugeCard.jsx ---
    doc.add_heading("3.4. Composant Jauge circulaire SVG (GaugeCard.jsx)", level=2)
    add_p(doc,
        "Composant réutilisable qui affiche une jauge circulaire SVG avec animation. "
        "La couleur change selon le type de capteur (vert=pH, bleu=EC, orange=température, "
        "turquoise=niveau d'eau) et passe en rouge si la valeur est hors de la plage optimale."
    )
    add_code_block(doc,
        read_file(os.path.join(BASE_DIR, "frontend", "src", "components", "GaugeCard.jsx")),
        "JavaScript (JSX)", "frontend/src/components/GaugeCard.jsx"
    )

    # --- 3.5 Control.jsx ---
    doc.add_heading("3.5. Contrôle du système (Control.jsx)", level=2)
    add_p(doc,
        "Écran de contrôle permettant de basculer entre mode automatique et manuel, "
        "de contrôler la pompe et l'éclairage via des interrupteurs, et de régler les "
        "consignes pH/EC avec des curseurs."
    )
    add_code_block(doc,
        read_file(os.path.join(BASE_DIR, "frontend", "src", "pages", "Control.jsx")),
        "JavaScript (JSX)", "frontend/src/pages/Control.jsx"
    )

    # --- 3.6 History.jsx ---
    doc.add_heading("3.6. Historique des capteurs (History.jsx)", level=2)
    add_p(doc,
        "Écran d'historique avec graphiques interactifs Chart.js montrant l'évolution "
        "des mesures sur les 60 dernières minutes. Inclut 4 onglets (pH, EC, Température, "
        "Niveau d'eau) et des statistiques (moyenne, minimum, maximum)."
    )
    add_code_block(doc,
        read_file(os.path.join(BASE_DIR, "frontend", "src", "pages", "History.jsx")),
        "JavaScript (JSX)", "frontend/src/pages/History.jsx"
    )

    # --- 3.7 Alerts.jsx ---
    doc.add_heading("3.7. Système d'alertes (Alerts.jsx)", level=2)
    add_p(doc,
        "Écran affichant la liste des alertes avec icônes par type (pH, EC, température, "
        "pompe, niveau d'eau), sévérité (warning/critical), horodatage et boutons d'acquittement."
    )
    add_code_block(doc,
        read_file(os.path.join(BASE_DIR, "frontend", "src", "pages", "Alerts.jsx")),
        "JavaScript (JSX)", "frontend/src/pages/Alerts.jsx"
    )

    # --- 3.8 Cultures.jsx ---
    doc.add_heading("3.8. Gestion des cultures (Cultures.jsx)", level=2)
    add_p(doc,
        "Écran affichant les 3 profils de cultures (Laitue, Basilic, Fraise) avec leurs "
        "paramètres optimaux. La sélection d'une culture ajuste automatiquement les seuils "
        "du système de contrôle."
    )
    add_code_block(doc,
        read_file(os.path.join(BASE_DIR, "frontend", "src", "pages", "Cultures.jsx")),
        "JavaScript (JSX)", "frontend/src/pages/Cultures.jsx"
    )

    # --- 3.9 Navbar.jsx ---
    doc.add_heading("3.9. Barre de navigation (Navbar.jsx)", level=2)
    add_p(doc,
        "Composant de navigation fixe en bas de l'écran avec 5 icônes correspondant aux "
        "5 écrans. Un badge rouge affiche le nombre d'alertes non acquittées."
    )
    add_code_block(doc,
        read_file(os.path.join(BASE_DIR, "frontend", "src", "components", "Navbar.jsx")),
        "JavaScript (JSX)", "frontend/src/components/Navbar.jsx"
    )

    doc.add_page_break()

    # =============================================
    # 4. CODE SOURCE BACKEND (Python)
    # =============================================
    doc.add_heading("4. Code source du Backend (Python Flask)", level=1)
    add_p(doc,
        "Le backend est développé en Python avec le framework Flask. "
        "Il expose une API REST de 12 endpoints et gère la simulation des capteurs, "
        "la logique de contrôle automatique et la base de données SQLite."
    )

    # --- 4.1 app.py ---
    doc.add_heading("4.1. Serveur API REST (app.py)", level=2)
    add_p(doc,
        "Fichier principal du serveur Flask. Il définit les 12 endpoints de l'API REST "
        "organisés en 4 catégories : capteurs, cultures, contrôle et alertes. "
        "Il lance également le thread de contrôle automatique et le simulateur de capteurs."
    )
    add_code_block(doc,
        read_file(os.path.join(BASE_DIR, "backend", "app.py")),
        "Python", "backend/app.py"
    )

    # --- 4.2 models.py ---
    doc.add_heading("4.2. Modèles de données SQLite (models.py)", level=2)
    add_p(doc,
        "Définition des 4 tables de la base de données SQLite et initialisation "
        "avec les données par défaut des 3 cultures du rapport PFE (Laitue, Basilic, Fraise)."
    )
    add_code_block(doc,
        read_file(os.path.join(BASE_DIR, "backend", "models.py")),
        "Python", "backend/models.py"
    )

    # --- 4.3 sensor_simulator.py ---
    doc.add_heading("4.3. Simulateur de capteurs ESP32 (sensor_simulator.py)", level=2)
    add_p(doc,
        "Simulateur réaliste des 4 capteurs connectés à l'ESP32 : capteur pH analogique, "
        "capteur EC DFR0300, capteur température DS18B20 et capteur niveau d'eau ST045. "
        "Les valeurs sont générées avec du bruit gaussien et une dérive naturelle, reproduisant "
        "fidèlement le comportement des capteurs physiques."
    )
    add_code_block(doc,
        read_file(os.path.join(BASE_DIR, "backend", "sensor_simulator.py")),
        "Python", "backend/sensor_simulator.py"
    )

    # --- 4.4 control_logic.py ---
    doc.add_heading("4.4. Logique de contrôle automatique (control_logic.py)", level=2)
    add_p(doc,
        "Implémentation de la logique de contrôle automatique conforme au diagramme de cas "
        "d'utilisation du rapport PFE. Vérifie périodiquement (toutes les 5 secondes) les "
        "valeurs des capteurs et déclenche des actions : arrêt de la pompe si pH ou EC "
        "hors plage, gestion de l'éclairage selon les heures de lumière de la culture active, "
        "et génération d'alertes avec mécanisme anti-spam."
    )
    add_code_block(doc,
        read_file(os.path.join(BASE_DIR, "backend", "control_logic.py")),
        "Python", "backend/control_logic.py"
    )

    doc.add_page_break()

    # =============================================
    # 5. CODE SOURCE CSS
    # =============================================
    doc.add_heading("5. Code source du Design (CSS3)", level=1)
    add_p(doc,
        "Le fichier CSS constitue le Design System complet de l'application. "
        "Il définit les variables de couleurs, les composants visuels, les animations "
        "et le layout responsive. Le thème Dark Mode utilise une palette de couleurs "
        "inspirée de la nature (vert, bleu, turquoise)."
    )

    doc.add_heading("5.1. Design System complet (index.css — 966 lignes)", level=2)
    add_p(doc,
        "Ce fichier contient l'intégralité du Design System : variables CSS, reset, "
        "layout, composants (jauges, cartes, toggles, sliders, alertes, graphiques), "
        "navbar, animations et media queries pour le responsive."
    )

    css_content = read_file(os.path.join(BASE_DIR, "frontend", "src", "index.css"))
    add_code_block(doc, css_content, "CSS3", "frontend/src/index.css")

    doc.add_page_break()

    # =============================================
    # RESUME FINAL
    # =============================================
    doc.add_heading("Récapitulatif", level=1)

    add_table(doc,
        ["Métrique", "Valeur"],
        [
            ["Nombre total de fichiers source", "12"],
            ["Lignes de code JavaScript (JSX)", "~750"],
            ["Lignes de code Python", "~844"],
            ["Lignes de code CSS", "966"],
            ["Total lignes de code", "~2 560"],
            ["Langage principal", "JavaScript (ES6+ / React.js / Node.js)"],
            ["Langage backend", "Python (Flask)"],
            ["Base de données", "SQLite (4 tables)"],
            ["Nombre d'endpoints API REST", "12"],
            ["Nombre d'écrans (pages)", "5"],
            ["Nombre de composants React", "7 (5 pages + 2 composants)"],
        ],
        caption="Tableau 5 : Récapitulatif du code source de l'application HydroNFT"
    )

    add_p(doc,
        "Ce document présente l'intégralité du code source de l'application HydroNFT, "
        "développée dans le cadre du Projet de Fin d'Études. Le langage principal est "
        "JavaScript (ES6+) utilisé avec le framework React.js pour le frontend et "
        "Node.js/Express pour le backend. Un backend alternatif en Python (Flask) "
        "a également été développé. L'application gère la surveillance en temps réel "
        "des capteurs (pH, EC, température, niveau d'eau), le contrôle automatique des "
        "actionneurs (pompe et éclairage) et le système d'alertes.",
        space_after=12
    )


def main():
    """Point d'entrée principal."""
    doc = Document()
    setup_styles(doc)

    # Marges du document
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    build_document(doc)
    doc.save(OUTPUT)

    print("=" * 60)
    print("  Document genere avec succes !")
    print(f"  Fichier : {OUTPUT}")
    print("=" * 60)
    print("\n  Contenu :")
    print("    - Langages de programmation (JavaScript, Python, HTML/CSS, SQL)")
    print("    - Architecture et structure du projet")
    print("    - Code source Frontend React.js (7 fichiers)")
    print("    - Code source Backend Python Flask (4 fichiers)")
    print("    - Code source Design CSS (1 fichier)")
    print("    - 5 tableaux recapitulatifs")
    print("    - Arborescence complete du projet")


if __name__ == "__main__":
    main()
