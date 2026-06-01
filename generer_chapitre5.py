# -*- coding: utf-8 -*-
"""
Generation du Chapitre 5 COMPLET du Rapport PFE de Hadil BOUDHRIOUWA
AVEC captures d'ecran reelles + diagrammes UML
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SCREENSHOT_DIR = os.path.join(BASE_DIR, "screenshots")
DIAGRAM_DIR = os.path.join(BASE_DIR, "diagrammes")
OUTPUT = os.path.join(BASE_DIR, "Chapitre5_Application_Mobile.docx")

SCREENSHOTS = {
    "splash": "01_splash_screen.png",
    "dashboard": "02_dashboard.png",
    "controle": "03_controle.png",
    "cultures": "04_cultures.png",
    "historique_ph": "05_historique_ph.png",
    "historique_ec": "06_historique_ec.png",
    "historique_temp": "07_historique_temp.png",
    "alertes": "08_alertes.png",
    "controle_manuel": "09_controle_manuel.png",
    "dashboard_final": "10_dashboard_final.png",
}

DIAGRAMS = {
    "use_case": "01_use_case.png",
    "seq_monitoring": "02_sequence_monitoring.png",
    "seq_controle": "03_sequence_controle.png",
    "database": "04_database_schema.png",
    "navigation": "05_navigation.png",
    "architecture": "06_architecture.png",
    "activite": "07_activite_controle.png",
}

fig_counter = [21]  # compteur global pour les figures


def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for level, size in [(1, 16), (2, 14), (3, 13)]:
        s = doc.styles[f'Heading {level}']
        s.font.name = 'Times New Roman'
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.space_before = Pt(18 if level == 1 else 14)
        s.paragraph_format.space_after = Pt(10)
        s.paragraph_format.line_spacing = 1.5


def add_p(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.25)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def next_fig():
    fig_counter[0] += 1
    return fig_counter[0]


def add_image(doc, folder, filename, caption, width=2.8):
    """Ajouter une image avec legende."""
    img_path = os.path.join(folder, filename)
    fig_num = next_fig()
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width))
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Fig. {fig_num} : {caption}")
    run2.bold = True
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(11)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    return fig_num


def add_diagram(doc, key, caption, width=5.5):
    """Ajouter un diagramme (large)."""
    return add_image(doc, DIAGRAM_DIR, DIAGRAMS[key], caption, width)


def add_screenshot(doc, key, caption, width=2.8):
    """Ajouter une capture d'ecran."""
    return add_image(doc, SCREENSHOT_DIR, SCREENSHOTS[key], caption, width)


def add_two_screenshots(doc, key1, key2, caption, width=2.2):
    """Deux captures cote a cote."""
    img1 = os.path.join(SCREENSHOT_DIR, SCREENSHOTS[key1])
    img2 = os.path.join(SCREENSHOT_DIR, SCREENSHOTS[key2])
    fig_num = next_fig()

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, img in enumerate([img1, img2]):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(img):
            run = p.add_run()
            run.add_picture(img, width=Inches(width))

    # Supprimer bordures
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = tbl.makeelement(qn('w:tblPr'), {})
        tbl.insert(0, tblPr)
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = borders.makeelement(qn(f'w:{bn}'), {
            qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'
        })
        borders.append(b)
    tblPr.append(borders)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Fig. {fig_num} : {caption}")
    run2.bold = True
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(11)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    return fig_num


def add_table(doc, headers, rows, caption=None, caption_num=None):
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(f"Tableau {caption_num} : {caption}")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = cell._element.get_or_add_tcPr()
        s = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'D9E2F3', qn('w:val'): 'clear'})
        shading.append(s)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def build_chapter(doc):
    # ═══════════════════════════════════════════
    # TITRE DU CHAPITRE
    # ═══════════════════════════════════════════
    title = doc.add_heading("Chapitre 5 : Conception et developpement de l'application mobile", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 1. INTRODUCTION ──
    doc.add_heading("1. Introduction :", level=2)
    add_p(doc,
        "Apres avoir concu et realise le systeme hydroponique NFT intelligent avec ses capteurs "
        "et son microcontroleur ESP32, il est essentiel de disposer d'une interface permettant "
        "de visualiser les donnees collectees et de controler le systeme a distance. "
        "Dans ce chapitre, nous presentons la conception et le developpement de l'application "
        "mobile HydroNFT, qui constitue l'interface homme-machine (IHM) de notre systeme. "
        "Cette application permet le monitoring en temps reel des parametres environnementaux, "
        "le controle a distance des actionneurs (pompe et eclairage), et la gestion des alertes."
    )

    # ═══════════════════════════════════════════
    # 2. ENVIRONNEMENT DE DEVELOPPEMENT
    # ═══════════════════════════════════════════
    doc.add_heading("2. Presentation de l'environnement de developpement :", level=2)
    add_p(doc,
        "Le developpement de l'application HydroNFT repose sur une architecture client-serveur "
        "moderne, utilisant des technologies web adaptees au deploiement sur appareils mobiles. "
        "L'approche Progressive Web App (PWA) a ete choisie pour permettre un acces universel "
        "depuis n'importe quel navigateur, sans necessiter d'installation depuis un store."
    )

    doc.add_heading("2.1. Technologies utilisees :", level=3)
    add_table(doc,
        ["Couche", "Technologie", "Version", "Role"],
        [
            ["Frontend", "HTML5 / CSS3 / JavaScript", "ES6+", "Interface utilisateur responsive"],
            ["Frontend", "React.js", "19.2", "Framework d'interface (composants)"],
            ["Frontend", "Vite", "8.0", "Outil de build et serveur de dev"],
            ["Frontend", "Chart.js", "4.5", "Visualisation graphique des donnees"],
            ["Backend", "Node.js + Express", "4.21", "Serveur API REST"],
            ["Base de donnees", "SQLite (sql.js)", "--", "Stockage des lectures et alertes"],
            ["Deploiement", "Render.com", "--", "Hebergement cloud du serveur"],
        ],
        caption="Technologies utilisees pour le developpement de l'application",
        caption_num=7
    )

    doc.add_heading("2.2. Architecture de l'application :", level=3)
    add_p(doc,
        "L'application suit une architecture client-serveur a trois couches (3-tier), "
        "separant clairement la presentation, la logique metier et les donnees :"
    )
    add_bullet(doc, "Couche Presentation : Interface utilisateur avec 5 ecrans principaux, "
               "design responsive mobile-first, navigation par barre inferieure.")
    add_bullet(doc, "Couche Logique Metier : Serveur Express.js (Node.js) exposant 12 endpoints "
               "API REST pour la lecture des capteurs, le controle des actionneurs et la gestion des alertes.")
    add_bullet(doc, "Couche Donnees : Base de donnees SQLite embarquee contenant 4 tables.")

    add_p(doc,
        "La figure suivante presente l'architecture de deploiement du systeme complet, "
        "montrant les interactions entre le systeme physique (ESP32 + capteurs), le serveur "
        "cloud et l'application cliente :")

    add_diagram(doc, "architecture", "Architecture de deploiement du systeme HydroNFT")

    doc.add_heading("2.3. Structure du projet :", level=3)
    add_table(doc,
        ["Dossier / Fichier", "Role"],
        [
            ["server.js", "Serveur backend Node.js (Express + simulation capteurs + API REST)"],
            ["index.html", "Page HTML principale (5 sections / ecrans)"],
            ["app.js", "Logique frontend JavaScript (navigation, jauges SVG, graphiques)"],
            ["styles.css", "Design System complet (Dark Mode Nature-Tech)"],
            ["backend/models.py", "Modeles de donnees SQLite et seed des 3 cultures"],
            ["backend/sensor_simulator.py", "Simulateur realiste des capteurs ESP32"],
            ["backend/control_logic.py", "Logique de controle automatique"],
        ],
        caption="Structure des fichiers du projet HydroNFT",
        caption_num=8
    )

    doc.add_heading("2.4. Schema de la base de donnees :", level=3)
    add_p(doc,
        "La base de donnees SQLite est composee de 4 tables qui stockent l'ensemble des "
        "donnees du systeme : les lectures des capteurs, les profils de cultures, l'etat "
        "du systeme et les alertes. La table system_state possede une cle etrangere vers "
        "culture_profiles pour identifier la culture active."
    )

    add_diagram(doc, "database", "Schema de la base de donnees SQLite (4 tables)")

    # ═══════════════════════════════════════════
    # 3. INTERFACE UTILISATEUR (UI/UX)
    # ═══════════════════════════════════════════
    doc.add_heading("3. Interface utilisateur (UI/UX) :", level=2)
    add_p(doc,
        "L'interface de l'application a ete concue selon les principes du Material Design "
        "et du mobile-first design, offrant une experience utilisateur intuitive et moderne. "
        "Le design adopte un theme sombre (Dark Mode) avec une palette de couleurs inspiree "
        "de la nature."
    )

    doc.add_heading("3.1. Design System :", level=3)
    add_table(doc,
        ["Element", "Valeur", "Description"],
        [
            ["Couleur primaire", "#2ECC71 (Vert)", "Bon fonctionnement"],
            ["Couleur secondaire", "#3498DB (Bleu)", "Informations, capteur EC"],
            ["Couleur tertiaire", "#1ABC9C (Turquoise)", "Accent, niveau d'eau"],
            ["Couleur danger", "#E74C3C (Rouge)", "Alertes critiques"],
            ["Couleur warning", "#F39C12 (Orange)", "Avertissements"],
            ["Fond principal", "#0a0e17", "Arriere-plan Dark Mode"],
            ["Police titres", "Outfit (700-800)", "Google Fonts"],
            ["Police corps", "Inter (400-600)", "Google Fonts"],
        ],
        caption="Palette de couleurs et typographie du Design System",
        caption_num=9
    )

    doc.add_heading("3.2. Diagramme de navigation :", level=3)
    add_p(doc,
        "La navigation entre les differents ecrans s'effectue via une barre de navigation "
        "fixe en bas de l'ecran (bottom navbar) contenant 5 icones. Le diagramme suivant "
        "illustre le flux de navigation entre les ecrans :"
    )
    add_diagram(doc, "navigation", "Diagramme de navigation entre les 5 ecrans")

    # ── 3.3 Splash Screen ──
    doc.add_heading("3.3. Ecran d'accueil (Splash Screen) :", level=3)
    add_p(doc,
        "Au lancement de l'application, un ecran d'accueil anime (Splash Screen) s'affiche "
        "pendant 1,5 seconde. Il presente le logo HydroNFT avec une animation de gradient "
        "vert-turquoise, accompagne d'un spinner de chargement. Cet ecran permet au serveur "
        "d'initialiser la connexion avec la base de donnees."
    )
    add_screenshot(doc, "splash", "Ecran d'accueil (Splash Screen) de l'application HydroNFT")

    # ── 3.4 Dashboard ──
    doc.add_heading("3.4. Tableau de bord (Dashboard) :", level=3)
    add_p(doc,
        "Le tableau de bord constitue l'ecran principal. Il offre une vue d'ensemble en temps "
        "reel de l'etat du systeme hydroponique, avec un rafraichissement automatique toutes "
        "les 3 secondes. Il se compose de trois sections :"
    )
    add_p(doc, "a) Grille de 4 jauges circulaires SVG (2x2) :", bold=True, space_after=3)
    add_bullet(doc, "pH (potentiel hydrogene) : plage 0-14, optimal 5.5-6.5, couleur verte")
    add_bullet(doc, "EC (conductivite electrique) : plage 0-5 mS/cm, optimal 0.8-2.2, couleur bleue")
    add_bullet(doc, "Temperature : plage 0-50 C, optimal 18-24 C, couleur orange")
    add_bullet(doc, "Niveau d'eau : plage 0-100%, seuil minimum 25%, couleur turquoise")

    add_p(doc, "b) Puces de statut : Pompe ON/OFF, Eclairage ON/OFF, Mode Auto/Manuel.", bold=True, space_after=3)
    add_p(doc, "c) Banniere d'alerte la plus recente (rouge=critique, orange=warning).", bold=True, space_after=3)

    add_screenshot(doc, "dashboard", "Tableau de bord avec les 4 jauges, statuts et banniere d'alerte")

    # ── 3.5 Capteurs en temps reel ──
    doc.add_heading("3.5. Affichage des donnees des capteurs en temps reel :", level=3)
    add_p(doc,
        "L'ecran Historique permet de visualiser l'evolution des mesures sous forme de "
        "graphiques interactifs generes avec Chart.js (version 4.5). Sous le graphique, "
        "trois cartes affichent les statistiques : Moyenne, Minimum et Maximum."
    )
    add_bullet(doc, "Courbe lissee avec remplissage semi-transparent sous la courbe")
    add_bullet(doc, "Donnees des 60 dernieres minutes (200 points max)")
    add_bullet(doc, "Rafraichissement automatique toutes les 10 secondes")
    add_bullet(doc, "4 onglets : pH, EC, Temperature, Niveau d'eau")

    add_two_screenshots(doc, "historique_ph", "historique_ec",
                        "Ecran Historique - Graphiques pH (gauche) et EC (droite)")

    # ── 3.6 Alertes ──
    doc.add_heading("3.6. Systeme d'alertes et notifications :", level=3)
    add_p(doc,
        "Le systeme d'alertes assure la surveillance proactive du systeme hydroponique :"
    )

    add_table(doc,
        ["Condition", "Type", "Severite", "Action"],
        [
            ["pH hors plage", "pH", "Warning/Critical", "Arret pompe"],
            ["EC hors plage", "EC", "Warning/Critical", "Arret pompe"],
            ["Niveau d'eau < 25%", "Niveau", "Warning", "Notification"],
            ["Niveau d'eau < 10%", "Niveau", "Critical", "Alerte critique"],
            ["Temperature hors plage", "Temp.", "Warning", "Notification"],
            ["Pompe arretee auto", "Pompe", "Critical", "Notification"],
        ],
        caption="Conditions de declenchement des alertes",
        caption_num=10
    )

    add_p(doc,
        "Un mecanisme anti-spam empeche les doublons (delai de 60 secondes). "
        "Un badge rouge dans la navbar indique le nombre d'alertes non acquittees."
    )
    add_screenshot(doc, "alertes", "Ecran Alertes avec liste des notifications et boutons d'acquittement")

    # ── 3.7 Controle pompe ──
    doc.add_heading("3.7. Controle a distance de la pompe et de l'eclairage :", level=3)
    add_p(doc,
        "L'ecran Controle permet de gerer les actionneurs et de configurer les seuils :"
    )
    add_bullet(doc, "Mode Automatique : le systeme gere pompe et eclairage selon les capteurs et seuils.")
    add_bullet(doc, "Mode Manuel : l'utilisateur controle directement chaque actionneur.")
    add_bullet(doc, "Pompe a eau (600 L/H) : interrupteur ON/OFF")
    add_bullet(doc, "Eclairage LED : interrupteur ON/OFF")
    add_bullet(doc, "4 curseurs pour regler les seuils pH min/max et EC min/max")

    add_two_screenshots(doc, "controle", "controle_manuel",
                        "Ecran Controle - Mode Automatique (gauche) et Mode Manuel (droite)")

    # ── 3.8 Cultures ──
    doc.add_heading("3.8. Gestion des cultures :", level=3)
    add_p(doc,
        "L'ecran Mes Cultures affiche les 3 profils pre-configures avec les parametres "
        "optimaux du Chapitre 3 :"
    )

    add_table(doc,
        ["Culture", "pH", "EC (mS/cm)", "Temp. (C)", "Humidite (%)", "Lumiere (h)"],
        [
            ["Laitue", "5.5 - 6.5", "0.8 - 1.4", "18 - 22", "50 - 70", "10 - 14"],
            ["Basilic", "5.5 - 6.5", "1.2 - 1.6", "18 - 24", "40 - 60", "12 - 16"],
            ["Fraise", "5.5 - 6.5", "1.8 - 2.2", "18 - 24", "60 - 75", "12 - 14"],
        ],
        caption="Parametres optimaux des cultures supportees",
        caption_num=11
    )

    add_p(doc,
        "La selection d'une culture ajuste automatiquement les seuils du systeme. "
        "La culture active est mise en evidence par une bordure verte lumineuse et un badge Active."
    )
    add_screenshot(doc, "cultures", "Ecran Cultures - Profils Laitue et Basilic avec parametres")

    # ═══════════════════════════════════════════
    # 4. DIAGRAMMES UML
    # ═══════════════════════════════════════════
    doc.add_heading("4. Modelisation UML de l'application :", level=2)
    add_p(doc,
        "La modelisation UML permet de formaliser les interactions entre l'utilisateur, "
        "l'application et le systeme physique. Nous presentons ici les principaux diagrammes "
        "couvrant les cas d'utilisation, les sequences d'interaction et la logique de controle."
    )

    # 4.1 Use Case
    doc.add_heading("4.1. Diagramme de cas d'utilisation :", level=3)
    add_p(doc,
        "Le diagramme de cas d'utilisation identifie les 7 fonctionnalites principales "
        "offertes a l'utilisateur par l'application mobile. L'acteur secondaire ESP32 "
        "intervient dans les cas impliquant la collecte de donnees capteurs."
    )
    add_diagram(doc, "use_case", "Diagramme de cas d'utilisation de l'application HydroNFT")

    # 4.2 Sequence monitoring
    doc.add_heading("4.2. Diagramme de sequence - Monitoring des capteurs :", level=3)
    add_p(doc,
        "Le diagramme de sequence suivant illustre le flux de communication pour le "
        "monitoring en temps reel. L'application interroge le serveur toutes les 3 secondes "
        "via l'API REST, qui retourne les donnees des capteurs depuis la base SQLite. "
        "Les jauges SVG sont mises a jour dynamiquement cote frontend."
    )
    add_diagram(doc, "seq_monitoring",
                "Diagramme de sequence - Monitoring des capteurs en temps reel")

    # 4.3 Sequence controle
    doc.add_heading("4.3. Diagramme de sequence - Controle de la pompe :", level=3)
    add_p(doc,
        "Ce diagramme illustre le flux lorsque l'utilisateur active ou desactive la pompe "
        "via l'interface. La commande transite du frontend vers le serveur API REST, qui "
        "met a jour l'etat dans la base de donnees et transmet la commande a l'ESP32."
    )
    add_diagram(doc, "seq_controle",
                "Diagramme de sequence - Controle a distance de la pompe")

    # 4.4 Activite
    doc.add_heading("4.4. Diagramme d'activite - Logique de controle automatique :", level=3)
    add_p(doc,
        "Le diagramme d'activite suivant formalise la logique de controle automatique "
        "executee toutes les 5 secondes par le serveur. Le systeme verifie sequentiellement "
        "le pH, l'EC, le niveau d'eau et la temperature, declenchant des alertes et arretant "
        "la pompe si necessaire."
    )
    add_diagram(doc, "activite",
                "Diagramme d'activite - Logique de controle automatique", width=4.0)

    # ═══════════════════════════════════════════
    # 5. COMMUNICATION
    # ═══════════════════════════════════════════
    doc.add_heading("5. Communication entre l'application et le systeme (Wi-Fi) :", level=2)
    add_p(doc,
        "La communication entre l'application et le systeme hydroponique est assuree par le "
        "protocole Wi-Fi integre dans l'ESP32 (2.4 GHz, 150 Mbit/s)."
    )

    doc.add_heading("5.1. API REST :", level=3)
    add_p(doc,
        "L'API expose 12 endpoints organises en 4 categories :"
    )
    add_table(doc,
        ["Categorie", "Methode", "Endpoint", "Description"],
        [
            ["Capteurs", "GET", "/api/sensors/current", "Valeurs actuelles des 4 capteurs"],
            ["Capteurs", "GET", "/api/sensors/history", "Historique + statistiques"],
            ["Cultures", "GET", "/api/cultures", "Liste des profils de cultures"],
            ["Cultures", "POST", "/api/cultures/:id/activate", "Activer une culture"],
            ["Controle", "POST", "/api/pump/toggle", "Demarrer/arreter la pompe"],
            ["Controle", "POST", "/api/lighting/toggle", "Allumer/eteindre l'eclairage"],
            ["Controle", "GET", "/api/system/status", "Etat complet du systeme"],
            ["Controle", "POST", "/api/system/mode", "Basculer Auto/Manuel"],
            ["Controle", "POST", "/api/settings/thresholds", "Regler les consignes"],
            ["Alertes", "GET", "/api/alerts", "Liste des alertes"],
            ["Alertes", "POST", "/api/alerts/:id/ack", "Acquitter une alerte"],
            ["Alertes", "POST", "/api/alerts/ack-all", "Acquitter toutes"],
        ],
        caption="Endpoints de l'API REST HydroNFT",
        caption_num=12
    )

    doc.add_heading("5.2. Protocole de rafraichissement :", level=3)
    add_table(doc,
        ["Ecran", "Intervalle", "Endpoint"],
        [
            ["Dashboard", "3 secondes", "/api/sensors/current"],
            ["Controle", "3 secondes", "/api/system/status"],
            ["Historique", "10 secondes", "/api/sensors/history"],
            ["Alertes", "5 secondes", "/api/alerts?limit=50"],
        ],
        caption="Intervalles de rafraichissement par ecran",
        caption_num=13
    )

    # ═══════════════════════════════════════════
    # 6. TESTS ET VALIDATION
    # ═══════════════════════════════════════════
    doc.add_heading("6. Tests et validation de l'application :", level=2)
    add_p(doc,
        "La validation a ete realisee en plusieurs phases, conformement au cycle en V "
        "(cf. Chapitre 1)."
    )

    doc.add_heading("6.1. Tests unitaires :", level=3)
    add_table(doc,
        ["Composant", "Test", "Resultat"],
        [
            ["Simulateur capteurs", "Plages pH (4.5-8.0), EC (0.3-5.0), temp (5-40 C)", "Conforme"],
            ["Logique de controle", "Arret pompe quand pH hors plage", "Conforme"],
            ["Logique de controle", "Eclairage auto selon heures culture", "Conforme"],
            ["Anti-spam alertes", "Pas de doublon dans les 60 sec.", "Conforme"],
            ["API REST", "Reponse 200 pour les 12 endpoints", "Conforme"],
            ["Activation culture", "Seuils mis a jour correctement", "Conforme"],
        ],
        caption="Resultats des tests unitaires",
        caption_num=14
    )

    doc.add_heading("6.2. Tests d'integration :", level=3)
    add_bullet(doc, "Capteur -> BDD : lectures enregistrees avec horodatage precis.")
    add_bullet(doc, "Capteur -> Alerte -> Frontend : alerte visible en moins de 5 secondes.")
    add_bullet(doc, "Frontend -> API -> Actionneur : commandes correctement relayees.")
    add_bullet(doc, "Culture -> Seuils -> Controle : ajustement automatique valide.")

    doc.add_heading("6.3. Tests de compatibilite :", level=3)
    add_table(doc,
        ["Appareil", "Resolution", "Resultat"],
        [
            ["Smartphone Android", "360 x 640 px", "Affichage correct"],
            ["Smartphone (grand)", "412 x 915 px", "Adaptation responsive"],
            ["Chrome Desktop", "1920 x 1080 px", "Centre 480px max-width"],
            ["Firefox", "1920 x 1080 px", "Compatible"],
            ["Edge", "1920 x 1080 px", "Compatible"],
        ],
        caption="Resultats des tests de compatibilite",
        caption_num=15
    )

    doc.add_heading("6.4. Validation des cas d'utilisation :", level=3)
    add_table(doc,
        ["Cas d'utilisation", "Scenario teste", "Resultat"],
        [
            ["Consulter capteurs", "Dashboard : 4 jauges avec valeurs actuelles", "Valide"],
            ["Demarrer/arreter pompe", "Toggle pompe : etat mis a jour", "Valide"],
            ["Controler eclairage", "Toggle eclairage : changement immediat", "Valide"],
            ["Regler consignes", "Curseurs pH/EC appliques", "Valide"],
            ["Changer de culture", "Selection Fraise : seuils ajustes", "Valide"],
            ["Consulter alertes", "Liste avec acquittement", "Valide"],
            ["Consulter historique", "Graphique + statistiques", "Valide"],
        ],
        caption="Validation des cas d'utilisation",
        caption_num=16
    )

    add_p(doc, "Les figures suivantes montrent l'application lors des tests :", bold=True)
    add_two_screenshots(doc, "dashboard", "dashboard_final",
                        "Dashboard au demarrage (gauche) et apres simulation (droite)")

    # ── 7. CONCLUSION ──
    doc.add_heading("7. Conclusion :", level=2)
    add_p(doc,
        "Dans ce chapitre, nous avons presente la conception et le developpement de "
        "l'application mobile HydroNFT. L'application, developpee avec des technologies "
        "web modernes (JavaScript, Node.js, Chart.js), offre une experience utilisateur "
        "complete et intuitive a travers 5 ecrans principaux."
    )
    add_p(doc,
        "Les diagrammes UML (cas d'utilisation, sequences, activite) formalisent les "
        "interactions entre l'utilisateur, l'application et le systeme physique. Le schema "
        "de la base de donnees et l'architecture de deploiement completent la documentation "
        "technique de la solution."
    )
    add_p(doc,
        "La communication Wi-Fi entre l'application et l'ESP32 garantit une connexion "
        "fiable, et les tests realises confirment le bon fonctionnement de l'ensemble de "
        "la chaine, de la collecte des donnees jusqu'a leur affichage dans l'interface."
    )


def main():
    doc = Document()
    setup_styles(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    build_chapter(doc)
    doc.save(OUTPUT)

    total_figs = fig_counter[0] - 21
    print(f"[OK] Chapitre 5 COMPLET genere !")
    print(f"    Fichier : {OUTPUT}")
    print(f"    Figures : {total_figs} (screenshots + diagrammes)")
    print(f"    Tableaux : 10")
    print(f"    Sections : 7 principales + sous-sections")
    print(f"    Contenu :")
    print(f"      - 10 captures d'ecran reelles de l'application")
    print(f"      - 7 diagrammes UML/techniques")
    print(f"      - 10 tableaux (techno, design, alertes, API, tests...)")


if __name__ == "__main__":
    main()
